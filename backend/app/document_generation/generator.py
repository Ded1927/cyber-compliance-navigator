import os
from io import BytesIO
from pathlib import Path

from docx import Document
from docxtpl import DocxTemplate
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter

from app.models.system_register_item import SystemRegisterItem


DOCX_MIME_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
XLSX_MIME_TYPE = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"

DEFAULT_TEMPLATE_PATH = (
    Path(__file__).resolve().parent / "templates" / "ciso_order_template.docx"
)


class DocumentGenerator:
    def __init__(self, ciso_order_template_path: str | Path | None = None):
        template_path = ciso_order_template_path or os.getenv("CISO_ORDER_TEMPLATE_PATH")
        self.ciso_order_template_path = (
            Path(template_path) if template_path else DEFAULT_TEMPLATE_PATH
        )

    def generate_ciso_order_docx(self, context: dict) -> BytesIO:
        self._ensure_ciso_order_template()
        template = DocxTemplate(self.ciso_order_template_path)
        template.render(self._normalize_ciso_context(context))

        output = BytesIO()
        template.save(output)
        output.seek(0)
        return output

    def generate_system_register_xlsx(
        self,
        systems: list[SystemRegisterItem],
    ) -> BytesIO:
        workbook = Workbook()
        worksheet = workbook.active
        worksheet.title = "System Register"

        headers = [
            "No.",
            "System name",
            "System type",
            "Information type",
            "OKII system",
            "Authorization status",
            "Evaluation status",
        ]
        worksheet.append(headers)

        for index, system in enumerate(systems, start=1):
            worksheet.append(
                [
                    index,
                    system.system_name,
                    system.system_type,
                    system.info_type,
                    "Yes" if system.is_okii_system else "No",
                    system.auth_status or "",
                    system.eval_status or "",
                ]
            )

        self._format_register_worksheet(worksheet)

        output = BytesIO()
        workbook.save(output)
        output.seek(0)
        return output

    def _normalize_ciso_context(self, context: dict) -> dict:
        organization_name = context.get("organization_name") or context.get("org_name") or ""
        ciso_full_name = context.get("ciso_full_name") or context.get("full_name") or ""
        department = context.get("department") or context.get("unit") or ""

        return {
            **context,
            "organization_name": organization_name,
            "ciso_full_name": ciso_full_name,
            "department": department,
        }

    def _ensure_ciso_order_template(self) -> None:
        if self.ciso_order_template_path.exists():
            return

        self.ciso_order_template_path.parent.mkdir(parents=True, exist_ok=True)

        document = Document()
        document.add_heading("Order on CISO Appointment", level=1)
        document.add_paragraph("Organization: {{ organization_name }}")
        document.add_paragraph("Appointed person: {{ ciso_full_name }}")
        document.add_paragraph("Department: {{ department }}")
        document.add_paragraph(
            "This order appoints {{ ciso_full_name }} as the person responsible "
            "for cybersecurity coordination in {{ organization_name }}."
        )
        document.add_paragraph("Basis: internal cybersecurity governance decision.")
        document.add_paragraph("Director: ____________________")
        document.save(self.ciso_order_template_path)

    def _format_register_worksheet(self, worksheet) -> None:
        header_fill = PatternFill("solid", fgColor="1F4E78")
        header_font = Font(color="FFFFFF", bold=True)
        thin_side = Side(style="thin", color="D9E2F3")
        border = Border(left=thin_side, right=thin_side, top=thin_side, bottom=thin_side)

        worksheet.freeze_panes = "A2"
        worksheet.auto_filter.ref = worksheet.dimensions

        for cell in worksheet[1]:
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            cell.border = border

        for row in worksheet.iter_rows(min_row=2):
            for cell in row:
                cell.alignment = Alignment(vertical="top", wrap_text=True)
                cell.border = border

        widths = [8, 32, 22, 24, 16, 24, 24]
        for index, width in enumerate(widths, start=1):
            worksheet.column_dimensions[get_column_letter(index)].width = width

        worksheet.row_dimensions[1].height = 28
